"""
resume_formatter.py
-------------------
Renders a TailoredResume Pydantic model into a well-formatted .docx file.

This is NOT a CrewAI tool (agents don't call it directly).
It's a utility called by the Flow after all crew runs are complete.

Design principles:
- Single-column, ATS-safe layout (no tables, text boxes, or columns)
- Consistent heading hierarchy
- Clean, readable typography using built-in Word styles
"""

from __future__ import annotations

from pathlib import Path

from docx.document import Document as DocumentClass
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from resume_builder.logger import get_logger
from resume_builder.models import TailoredResume

logger = get_logger(__name__)


# ─────────────────────────────────────────────
# Colour palette (ATS-safe, professional)
# ─────────────────────────────────────────────
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)  # near-black navy
ACCENT_COLOR = RGBColor(0x16, 0x47, 0x7D)  # muted blue for section rules
BODY_COLOR = RGBColor(0x2D, 0x2D, 0x2D)  # dark grey body text


class ResumeFormatterTool:
    """
    Converts a TailoredResume into a production-ready .docx file.

    Usage:
        formatter = ResumeFormatterTool()
        output_path = formatter.generate(resume, output_dir=Path("./outputs"))
    """

    def generate(
        self,
        resume: TailoredResume,
        output_dir: Path = Path("./outputs"),
    ) -> Path:
        """
        Generate the .docx file and write it to output_dir.
        Returns the path to the created file.
        """
        logger.info("Generating .docx for %s at %s", resume.job_title, resume.company)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / resume.output_filename()

        doc = Document()
        self._set_page_margins(doc)
        self._set_default_font(doc)

        self._write_header(doc, resume)
        self._write_summary(doc, resume)
        self._write_experience(doc, resume)
        if resume.projects:
            self._write_projects(doc, resume)
        self._write_skills(doc, resume)
        self._write_education(doc, resume)
        if resume.certifications:
            self._write_certifications(doc, resume)

        doc.save(str(output_path))
        logger.info("Saved .docx to %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------

    def _set_page_margins(self, doc: DocumentClass) -> None:
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

    def _set_default_font(self, doc: DocumentClass) -> None:
        style = doc.styles["Normal"]
        font = style.font  # type: ignore [reportAttributeAccessIssue]
        font.name = "Calibri"
        font.size = Pt(10.5)
        font.color.rgb = BODY_COLOR

    # ------------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------------

    def _write_header(self, doc: DocumentClass, resume: TailoredResume) -> None:
        c = resume.contact

        # Name — large, prominent
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(c.name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = HEADING_COLOR

        # Contact line
        contact_parts = filter(
            None, [c.email, c.phone, c.location, c.linkedin, c.github, c.portfolio]
        )
        contact_line = "  |  ".join(contact_parts)
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        self._add_horizontal_rule(doc)

    def _write_summary(self, doc: DocumentClass, resume: TailoredResume) -> None:
        self._add_section_heading(doc, "Professional Summary")
        doc.add_paragraph(resume.professional_summary)

    def _write_experience(self, doc: DocumentClass, resume: TailoredResume) -> None:
        self._add_section_heading(doc, "Professional Experience")

        for entry in resume.experience:
            # Role + company line
            role_para = doc.add_paragraph()
            role_run = role_para.add_run(f"{entry.role}")
            role_run.bold = True
            role_run.font.size = Pt(11)
            role_run.font.color.rgb = HEADING_COLOR

            role_para.add_run(f"  —  {entry.company}")

            # Dates line
            date_parts = f"{entry.start_date} – {entry.end_date}"
            if entry.location:
                date_parts += f"  ·  {entry.location}"
            date_para = doc.add_paragraph(date_parts)
            for run in date_para.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.italic = True

            # Bullet points
            for bullet in entry.bullets:
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_para.add_run(bullet)

    def _write_projects(self, doc: DocumentClass, resume: TailoredResume) -> None:
        self._add_section_heading(doc, "Projects")

        for proj in resume.projects:
            # Project name line
            proj_para = doc.add_paragraph()
            role_run = proj_para.add_run(f"{proj.repo_name}")
            role_run.bold = True
            role_run.font.size = Pt(11)
            role_run.font.color.rgb = HEADING_COLOR

            # Description line
            desc_para = doc.add_paragraph(proj.description)
            for run in desc_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = BODY_COLOR
                run.italic = True

            # Tech stack line
            tech_line = "  ·  ".join(proj.tech_stack[:8])
            if tech_line:
                tech_para = doc.add_paragraph(f"Tech: {tech_line}")
                for run in tech_para.runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Architecture bullet
            if proj.architecture:
                bullet_para = doc.add_paragraph(style="List Bullet")
                bullet_para.add_run(proj.architecture)

    def _write_skills(self, doc: DocumentClass, resume: TailoredResume) -> None:
        self._add_section_heading(doc, "Skills")
        # Single paragraph, skills separated by  ·
        skills_text = "  ·  ".join(resume.skills)
        doc.add_paragraph(skills_text)

    def _write_education(self, doc: DocumentClass, resume: TailoredResume) -> None:
        self._add_section_heading(doc, "Education")
        for entry in resume.education:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(f"{entry.degree}")
            degree_run.bold = True
            if entry.field_of_study:
                edu_para.add_run(f" in {entry.field_of_study}")
            edu_para.add_run(f"  —  {entry.institution}")
            if entry.start_date:
                edu_para.add_run(f"  ({entry.start_date})")
            if entry.end_date:
                edu_para.add_run(f"  ({entry.end_date})")
            if entry.honours:
                doc.add_paragraph(entry.honours)

    def _write_certifications(self, doc: DocumentClass, resume: TailoredResume) -> None:
        self._add_section_heading(doc, "Certifications")
        for cert in resume.certifications:
            cert_para = doc.add_paragraph(style="List Bullet")
            cert_para.add_run(cert)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _add_section_heading(self, doc: DocumentClass, title: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = ACCENT_COLOR
        run.font.name = "Calibri"

        # Bottom border on the paragraph (acts as a section divider)
        self._add_paragraph_border(para, side="bottom")

    def _add_paragraph_border(
        self,
        paragraph,
        side: str = "bottom",
        color: str = "164779",
        size: int = 6,
    ) -> None:
        """Add a thin border to one side of a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "1")
        border.set(qn("w:color"), color)
        pBdr.append(border)
        pPr.append(pBdr)

    def _add_horizontal_rule(self, doc: DocumentClass) -> None:
        """Add a full-width horizontal rule paragraph."""
        para = doc.add_paragraph()
        self._add_paragraph_border(para, side="bottom", color="AAAAAA", size=4)
        para.paragraph_format.space_after = Pt(4)
