"""
test_resume_formatter.py
------------------------
Tests for ResumeFormatter — .docx generation from TailoredResume model.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from resume_builder.models import (
    AwardEntry,
    ContactInfo,
    EducationEntry,
    InternationalExperienceEntry,
    PublicationEntry,
    TailoredExperienceEntry,
    TailoredResume,
    WorkshopEntry,
)
from resume_builder.processors.formatter import ResumeFormatter


class TestResumeFormatter:
    def setup_method(self) -> None:
        self.formatter = ResumeFormatter()

    def _make_resume(self) -> TailoredResume:
        return TailoredResume(
            contact=ContactInfo(
                name="Alice Smith",
                email="alice@example.com",
                phone="+1-555-0100",
            ),
            professional_summary="Backend engineer with 5+ years of experience.",
            experience=[
                TailoredExperienceEntry(
                    company="TechCorp",
                    role="Senior Engineer",
                    start_date="Jan 2022",
                    end_date="Present",
                    location="Milan",
                    bullets=[
                        "Built microservices",
                        "Led migration to Go",
                    ],
                ),
            ],
            skills=["Python", "Go", "Docker"],
            education=[
                EducationEntry(
                    institution="Politecnico di Milano",
                    degree="Master's",
                    field_of_study="CS",
                    start_date="2016",
                    end_date="2020",
                ),
            ],
            company="CloudScale",
            job_title="Platform Engineer",
            confidence_score=85,
            tailoring_notes="Tailored for platform role.",
            session_id=1,
            ats_keyword_coverage=["python", "go", "docker"],
        )

    def test_generates_docx(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        assert path.exists()
        assert path.suffix == ".docx"
        assert path.name == "resume_CloudScale_Platform_Engineer.docx"

    def test_docx_is_readable(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        full_text = " ".join(paragraphs)
        assert "Alice Smith" in full_text
        # The experience company (TechCorp) appears in the body
        assert "TechCorp" in full_text
        assert "Platform Engineer" not in full_text  # target job title, not shown
        assert "Backend engineer" in full_text

    def test_output_dir_created(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        nested = tmp_dir / "sub" / "dir"
        path = self.formatter.generate(resume, output_dir=nested)

        assert path.exists()
        assert path.parent == nested

    def test_with_certifications(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        resume.certifications = ["AWS Solutions Architect", "CKA"]
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "CERTIFICATIONS" in full_text.upper()
        assert "AWS Solutions Architect" in full_text

    def test_with_publications(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        resume.publications = [
            PublicationEntry(
                title="Attention Is All You Need",
                venue="NeurIPS 2017",
                date="Jun 2017",
                publisher="Curran Associates",
                link="https://arxiv.org/abs/1706.03762",
            )
        ]
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "PUBLICATIONS" in full_text.upper()
        assert "Attention Is All You Need" in full_text
        assert "NeurIPS 2017" in full_text
        assert "arxiv" in full_text

    def test_with_workshops(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        resume.workshops = [
            WorkshopEntry(
                title="MLOps Best Practices",
                date="Mar 2023",
                place="Milan, Italy",
            )
        ]
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "WORKSHOPS" in full_text.upper()
        assert "MLOps Best Practices" in full_text

    def test_with_awards(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        resume.awards = [
            AwardEntry(
                title="Best Paper Award",
                organization="IEEE",
                date="Dec 2022",
            )
        ]
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "AWARDS" in full_text.upper()
        assert "Best Paper Award" in full_text
        assert "IEEE" in full_text

    def test_with_international_experiences(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        resume.international_experiences = [
            InternationalExperienceEntry(
                place="Tokyo, Japan",
                date="Sep 2021 – Jun 2022",
                description="Exchange semester at University of Tokyo",
            )
        ]
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "INTERNATIONAL EXPERIENCES" in full_text.upper()
        assert "Tokyo, Japan" in full_text
        assert "Exchange semester" in full_text

    def test_all_sections_together(self, tmp_dir: Path) -> None:
        resume = self._make_resume()
        resume.publications = [
            PublicationEntry(
                title="Test Paper",
                venue="ICML 2023",
                date="Jul 2023",
            )
        ]
        resume.workshops = [
            WorkshopEntry(title="K8s Workshop", date="2022", place="Berlin")
        ]
        resume.awards = [
            AwardEntry(title="Innovation Award", organization="ACM", date="2021")
        ]
        resume.international_experiences = [
            InternationalExperienceEntry(
                place="Paris, France",
                date="2020",
                description="Research internship at INRIA",
            )
        ]
        path = self.formatter.generate(resume, output_dir=tmp_dir)

        doc = Document(str(path))
        full_text = " ".join(p.text for p in doc.paragraphs)
        for section in ["PUBLICATIONS", "WORKSHOPS", "AWARDS", "INTERNATIONAL EXPERIENCES"]:
            assert section in full_text.upper()
